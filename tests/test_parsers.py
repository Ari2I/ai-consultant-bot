"""
Тесты для documents.parsers.

Для DOCX и XLSX создаются реальные временные файлы (эти библиотеки
уже являются зависимостями проекта). Для PDF используется мок
pdfplumber.open — создание реального PDF потребовало бы отдельной
библиотеки только ради тестов, без реальной необходимости для кода.
"""

import os
import tempfile
import unittest
from unittest.mock import MagicMock, patch

import docx
import openpyxl

from documents.parsers import (
    DocumentParsingError,
    parse_docx,
    parse_document,
    parse_md,
    parse_pdf,
    parse_txt,
    parse_xlsx,
)


class ParseTxtTest(unittest.TestCase):
    def test_reads_utf8_file(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "about.txt")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("Магазин основан в 2020 году.")

            text = parse_txt(file_path)

            self.assertEqual(text, "Магазин основан в 2020 году.")

    def test_reads_cp1251_file_as_fallback(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "legacy.txt")
            with open(file_path, "w", encoding="cp1251") as f:
                f.write("Старый файл в кодировке Windows-1251.")

            text = parse_txt(file_path)

            self.assertEqual(text, "Старый файл в кодировке Windows-1251.")

    def test_missing_file_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_txt("/nonexistent/path/file.txt")

    def test_empty_file_raises(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "empty.txt")
            open(file_path, "w", encoding="utf-8").close()
            with self.assertRaises(DocumentParsingError):
                parse_txt(file_path)


class ParseMdTest(unittest.TestCase):
    def test_reads_markdown_file_as_is(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "faq.md")
            content = (
                "# Часто задаваемые вопросы\n\n"
                "## Режим работы\n\n"
                "Магазин работает с 10:00 до 20:00 без выходных."
            )
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

            text = parse_md(file_path)

            self.assertEqual(text, content)

    def test_reads_cp1251_markdown_as_fallback(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "legacy.md")
            with open(file_path, "w", encoding="cp1251") as f:
                f.write("# Старый файл\n\nВ кодировке Windows-1251.")

            text = parse_md(file_path)

            self.assertIn("Старый файл", text)

    def test_missing_file_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_md("/nonexistent/path/file.md")

    def test_empty_file_raises(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "empty.md")
            open(file_path, "w", encoding="utf-8").close()
            with self.assertRaises(DocumentParsingError):
                parse_md(file_path)


class ParseDocxTest(unittest.TestCase):
    def test_extracts_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "faq.docx")
            document = docx.Document()
            document.add_paragraph("Как оформить возврат?")
            document.add_paragraph("В течение 14 дней с чеком.")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "USD"
            table.rows[0].cells[1].text = "90.5"
            document.save(file_path)

            text = parse_docx(file_path)

            self.assertIn("Как оформить возврат?", text)
            self.assertIn("В течение 14 дней с чеком.", text)
            self.assertIn("USD", text)
            self.assertIn("90.5", text)

    def test_missing_file_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_docx("/nonexistent/path/file.docx")

    def test_empty_document_raises(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "empty.docx")
            docx.Document().save(file_path)
            with self.assertRaises(DocumentParsingError):
                parse_docx(file_path)


class ParseXlsxTest(unittest.TestCase):
    def test_extracts_rows_as_readable_text(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "price_list.xlsx")
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.append(["Товар", "Цена"])
            sheet.append(["Футболка", 1500])
            sheet.append(["Джинсы", 3200])
            workbook.save(file_path)

            text = parse_xlsx(file_path)

            self.assertIn("Товар", text)
            self.assertIn("Футболка", text)
            self.assertIn("1500", text)
            self.assertIn("Джинсы", text)

    def test_missing_file_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_xlsx("/nonexistent/path/file.xlsx")

    def test_empty_workbook_raises(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            file_path = os.path.join(tmp_dir, "empty.xlsx")
            openpyxl.Workbook().save(file_path)
            with self.assertRaises(DocumentParsingError):
                parse_xlsx(file_path)


class ParsePdfTest(unittest.TestCase):
    @patch("documents.parsers.pdfplumber.open")
    def test_extracts_text_from_all_pages(self, mock_open):
        page1 = MagicMock()
        page1.extract_text.return_value = "Страница 1"
        page2 = MagicMock()
        page2.extract_text.return_value = "Страница 2"
        mock_pdf = MagicMock()
        mock_pdf.pages = [page1, page2]
        mock_open.return_value.__enter__.return_value = mock_pdf

        text = parse_pdf("fake.pdf")

        self.assertIn("Страница 1", text)
        self.assertIn("Страница 2", text)

    @patch("documents.parsers.pdfplumber.open")
    def test_raises_when_no_extractable_text(self, mock_open):
        page = MagicMock()
        page.extract_text.return_value = None
        mock_pdf = MagicMock()
        mock_pdf.pages = [page]
        mock_open.return_value.__enter__.return_value = mock_pdf

        with self.assertRaises(DocumentParsingError):
            parse_pdf("fake.pdf")

    @patch("documents.parsers.pdfplumber.open")
    def test_open_error_raises_parsing_error(self, mock_open):
        mock_open.side_effect = OSError("файл повреждён")

        with self.assertRaises(DocumentParsingError):
            parse_pdf("fake.pdf")


class ParseDocumentDispatchTest(unittest.TestCase):
    @patch("documents.parsers.parse_txt")
    def test_dispatches_txt_by_extension(self, mock_parse_txt):
        mock_parse_txt.return_value = "текст"

        text, fmt = parse_document("about.txt")

        self.assertEqual(text, "текст")
        self.assertEqual(fmt, "txt")
        mock_parse_txt.assert_called_once_with("about.txt")

    @patch("documents.parsers.parse_md")
    def test_dispatches_md_by_extension(self, mock_parse_md):
        mock_parse_md.return_value = "текст"

        text, fmt = parse_document("faq.md")

        self.assertEqual(text, "текст")
        self.assertEqual(fmt, "md")
        mock_parse_md.assert_called_once_with("faq.md")

    @patch("documents.parsers.parse_xlsx")
    def test_dispatches_xlsx_case_insensitively(self, mock_parse_xlsx):
        mock_parse_xlsx.return_value = "текст"

        text, fmt = parse_document("price_list.XLSX")

        self.assertEqual(fmt, "xlsx")
        mock_parse_xlsx.assert_called_once()

    def test_unsupported_extension_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_document("photo.png")

    def test_missing_extension_raises(self):
        with self.assertRaises(DocumentParsingError):
            parse_document("file_without_extension")


if __name__ == "__main__":
    unittest.main()
