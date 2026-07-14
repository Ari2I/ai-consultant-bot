"""
Парсинг документов базы знаний: приведение PDF/DOCX/XLSX/TXT/MD к
обычному тексту, пригодному для дальнейшего чанкинга и вычисления
эмбеддингов.

Каждый парсер поднимает DocumentParsingError с понятным сообщением
при любой ошибке чтения файла или если из документа не удалось
извлечь ни одного символа текста (например, скан без слоя текста в
PDF, или пустой файл) — вызывающий код (обработчик /kb_add) должен
уметь показать эту ошибку админу, а не падать с трудночитаемым
исключением библиотеки.
"""

from __future__ import annotations

from pathlib import Path
from typing import Tuple

import docx
import openpyxl
import pdfplumber

# Кодировки, которые пробуются по очереди при чтении .txt: сначала
# UTF-8 (стандарт), затем CP1251 (частая кодировка старых русских
# текстовых файлов, сохранённых в Windows).
_TXT_FALLBACK_ENCODINGS = ("utf-8", "cp1251")


class DocumentParsingError(Exception):
    """Ошибка при чтении или разборе документа базы знаний."""


def parse_pdf(file_path: str) -> str:
    """
    Извлекает текст из PDF-файла постранично.

    Исключения:
        DocumentParsingError: файл повреждён/недоступен, либо ни на
            одной странице не нашлось извлекаемого текста (например,
            PDF — это скан изображений без текстового слоя).
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            pages_text = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:  # noqa: BLE001 — библиотека кидает разные типы
        raise DocumentParsingError(
            f"Не удалось прочитать PDF-файл {file_path}: {exc}"
        ) from exc

    text = "\n\n".join(
        page_text.strip() for page_text in pages_text if page_text.strip()
    )
    if not text.strip():
        raise DocumentParsingError(
            f"В PDF-файле {file_path} не найдено извлекаемого текста. "
            "Если это скан, потребуется OCR — в текущей версии не "
            "реализовано."
        )
    return text


def parse_docx(file_path: str) -> str:
    """
    Извлекает текст из Word-документа: обычные абзацы и таблицы.

    Ячейки таблиц объединяются в одну строку через " | " — этого
    достаточно для FAQ/прайс-листов, оформленных таблицей в Word.

    Исключения:
        DocumentParsingError: файл повреждён/недоступен либо не
            содержит ни одного непустого абзаца или ячейки таблицы.
    """
    try:
        document = docx.Document(file_path)
    except Exception as exc:  # noqa: BLE001
        raise DocumentParsingError(
            f"Не удалось прочитать DOCX-файл {file_path}: {exc}"
        ) from exc

    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cells_text = [
                cell.text.strip() for cell in row.cells if cell.text.strip()
            ]
            if cells_text:
                parts.append(" | ".join(cells_text))

    text = "\n\n".join(parts)
    if not text.strip():
        raise DocumentParsingError(
            f"DOCX-файл {file_path} не содержит текста (пустой документ)"
        )
    return text


def parse_xlsx(file_path: str) -> str:
    """
    Извлекает данные из Excel-файла в виде читаемого текста.

    Первая непустая строка каждого листа считается заголовком; для
    последующих строк с тем же количеством ячеек формируется строка
    вида "Колонка1: значение1, Колонка2: значение2" — это даёт модели
    больше контекста, чем просто перечисление значений через запятую.

    Исключения:
        DocumentParsingError: файл повреждён/недоступен либо во всех
            листах нет ни одной непустой строки.
    """
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise DocumentParsingError(
            f"Не удалось прочитать XLSX-файл {file_path}: {exc}"
        ) from exc

    parts = []
    for sheet in workbook.worksheets:
        header: list[str] | None = None
        for row in sheet.iter_rows(values_only=True):
            cells = [
                str(cell).strip()
                for cell in row
                if cell is not None and str(cell).strip()
            ]
            if not cells:
                continue
            if header is None:
                header = cells
                parts.append(" | ".join(cells))
                continue
            if len(cells) == len(header):
                parts.append(
                    ", ".join(f"{h}: {v}" for h, v in zip(header, cells))
                )
            else:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise DocumentParsingError(
            f"XLSX-файл {file_path} не содержит данных (пустая книга)"
        )
    return text


def _read_text_file(file_path: str) -> str:
    """
    Читает обычный текстовый файл, пробуя UTF-8, затем CP1251.

    Общая логика для форматов, где содержимое файла — это уже
    готовый текст без бинарной структуры (.txt, .md) — Markdown не
    требует отдельного разбора синтаксиса: заголовки, списки и
    выделения текста прекрасно понимает сама модель при генерации
    ответа, а лишний слой очистки разметки рискует случайно
    исказить смысл (например, потерять заголовок как часть текста).

    Исключения:
        DocumentParsingError: файл не найден/недоступен, не удалось
            определить кодировку, либо файл пустой.
    """
    path = Path(file_path)

    text = None
    last_error: Exception | None = None
    for encoding in _TXT_FALLBACK_ENCODINGS:
        try:
            text = path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError as exc:
            last_error = exc
            continue
        except OSError as exc:
            raise DocumentParsingError(
                f"Не удалось открыть файл {file_path}: {exc}"
            ) from exc

    if text is None:
        raise DocumentParsingError(
            f"Не удалось определить кодировку файла {file_path} "
            f"(пробовались: {', '.join(_TXT_FALLBACK_ENCODINGS)}): "
            f"{last_error}"
        )

    if not text.strip():
        raise DocumentParsingError(f"Файл {file_path} пуст")
    return text


def parse_txt(file_path: str) -> str:
    """Читает обычный текстовый файл (.txt). См. _read_text_file."""
    return _read_text_file(file_path)


def parse_md(file_path: str) -> str:
    """Читает Markdown-файл (.md) как обычный текст. См. _read_text_file."""
    return _read_text_file(file_path)


# Поддерживаемые расширения — используется только для сообщения об
# ошибке; сам выбор парсера ниже сделан явными вызовами (а не через
# словарь функций), чтобы подмена parse_txt/parse_docx/... в тестах
# через unittest.mock.patch срабатывала предсказуемо: вызов по имени
# всегда идёт через пространство имён модуля заново, в отличие от
# заранее сохранённой в словаре ссылки на функцию.
_SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".txt", ".md")


def parse_document(file_path: str) -> Tuple[str, str]:
    """
    Извлекает текст из документа, автоматически определяя формат по
    расширению файла.

    Возвращает:
        Кортеж (текст документа, формат без точки, например "docx").

    Исключения:
        DocumentParsingError: расширение не поддерживается, либо сам
            парсер для этого формата поднял ошибку.
    """
    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        text = parse_pdf(file_path)
    elif extension == ".docx":
        text = parse_docx(file_path)
    elif extension == ".xlsx":
        text = parse_xlsx(file_path)
    elif extension == ".txt":
        text = parse_txt(file_path)
    elif extension == ".md":
        text = parse_md(file_path)
    else:
        raise DocumentParsingError(
            f"Неподдерживаемый формат файла: "
            f"{extension or '(без расширения)'}. "
            f"Поддерживаются: {', '.join(_SUPPORTED_EXTENSIONS)}"
        )

    return text, extension.lstrip(".")
