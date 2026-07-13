"""
Скрипт генерации демо-документов базы знаний для вымышленного
магазина одежды «ТриНитки» — используется для демонстрации проекта
в портфолио (реальный заказчик загружает свои документы через
/kb_add, эти файлы нужны только для примера).

Запуск (создаёт/перезаписывает файлы в этой же директории):
    python demo_data/generate_demo_documents.py

Зависимости: python-docx и openpyxl уже входят в requirements.txt
проекта. Для генерации демо-PDF дополнительно требуется reportlab —
он НЕ входит в requirements.txt бота (не нужен для его работы,
только для регенерации этого демо-файла):
    pip install reportlab
"""

from __future__ import annotations

from pathlib import Path

import docx
import openpyxl
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUTPUT_DIR = Path(__file__).parent

# Встроенные шрифты reportlab (Helvetica, Times-Roman и т.д.) не
# поддерживают кириллицу — без регистрации отдельного TTF-шрифта
# русский текст в PDF превращается в нечитаемую кашу как при
# визуальном просмотре, так и (что важнее для этого проекта) при
# извлечении текста через pdfplumber. DejaVu Sans — свободный шрифт с
# полной поддержкой кириллицы, обычно уже установлен в Linux-
# дистрибутивах (пакет fonts-dejavu-core); при отсутствии в системе
# нужно установить его перед запуском скрипта.
_DEJAVU_PATHS = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
)


def _register_cyrillic_font() -> str:
    for font_path in _DEJAVU_PATHS:
        if Path(font_path).exists():
            pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))
            return "DejaVuSans"
    raise FileNotFoundError(
        "Не найден шрифт DejaVu Sans с поддержкой кириллицы. "
        "Установите пакет 'fonts-dejavu-core' (Debian/Ubuntu: "
        "sudo apt install fonts-dejavu-core) и запустите скрипт снова."
    )


def generate_about_txt() -> None:
    content = (
        "Магазин одежды «ТриНитки»\n\n"
        "«ТриНитки» — небольшой магазин повседневной одежды в Москве, "
        "работающий с 2019 года. Мы продаём базовую одежду для мужчин "
        "и женщин: футболки, худи, джинсы, лёгкие куртки — из плотного "
        "хлопка и переработанных материалов.\n\n"
        "Наша миссия — сделать качественную базовую одежду доступной "
        "и по разумной цене, без переплаты за бренд.\n\n"
        "Магазин расположен по адресу: Москва, ул. Тверская, д. 10 "
        "(шоурум работает только по предварительной записи). "
        "Интернет-магазин осуществляет доставку по всей России.\n\n"
        "Реквизиты для юридических лиц предоставляются по запросу "
        "через менеджера."
    )
    (OUTPUT_DIR / "about.txt").write_text(content, encoding="utf-8")


def generate_faq_docx() -> None:
    document = docx.Document()
    document.add_heading("Часто задаваемые вопросы — «ТриНитки»", level=1)

    faq_pairs = [
        (
            "Какой у вас режим работы шоурума?",
            "Шоурум в Москве работает по предварительной записи, с "
            "понедельника по пятницу с 11:00 до 19:00. Интернет-"
            "магазин принимает заказы круглосуточно.",
        ),
        (
            "Как оформить возврат или обмен?",
            "Возврат и обмен возможны в течение 14 дней с момента "
            "получения заказа, при сохранении бирок и товарного вида. "
            "Деньги возвращаются на карту в течение 5 рабочих дней "
            "после получения возврата магазином.",
        ),
        (
            "Какие способы оплаты вы принимаете?",
            "Оплата картой на сайте, а также наличными или картой при "
            "получении — для заказов по Москве.",
        ),
        (
            "Как узнать свой размер?",
            "На каждой странице товара есть таблица размеров с "
            "замерами в сантиметрах. Если сомневаетесь — напишите нам "
            "в этом чате, поможем подобрать размер.",
        ),
        (
            "Есть ли у вас программа лояльности?",
            "Да. При регистрации в личном кабинете вы получаете "
            "скидку 5% на второй заказ, далее действует накопительная "
            "система баллов за покупки.",
        ),
        (
            "Можно ли оформить заказ по телефону?",
            "Да, вы можете написать нам в этом чате-боте или "
            "позвонить менеджеру — контакты указаны на сайте магазина.",
        ),
    ]

    for question, answer in faq_pairs:
        document.add_paragraph(question, style="Heading 2")
        document.add_paragraph(answer)

    document.save(OUTPUT_DIR / "faq.docx")


def generate_price_list_xlsx() -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Прайс-лист"
    sheet.append(["Категория", "Товар", "Цена, руб.", "Размеры"])

    rows = [
        ("Футболки", "Футболка базовая хлопковая", 1500, "XS-XXL"),
        ("Футболки", "Футболка оверсайз", 1800, "S-XL"),
        ("Худи", "Худи на молнии", 3500, "S-XXL"),
        ("Худи", "Худи с капюшоном", 3200, "XS-XXL"),
        ("Джинсы", "Джинсы прямого кроя", 4200, "26-34"),
        ("Джинсы", "Джинсы slim", 3900, "26-32"),
        ("Куртки", "Лёгкая куртка-ветровка", 5500, "S-XL"),
        ("Аксессуары", "Шапка вязаная", 1200, "one size"),
    ]
    for row in rows:
        sheet.append(list(row))

    workbook.save(OUTPUT_DIR / "price_list.xlsx")


def generate_shipping_returns_pdf() -> None:
    font_name = _register_cyrillic_font()
    doc = SimpleDocTemplate(
        str(OUTPUT_DIR / "shipping_returns.pdf"), pagesize=A4
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.fontName = font_name
    body_style = styles["BodyText"]
    body_style.fontName = font_name

    story = [
        Paragraph(
            "Условия доставки и возврата — «ТриНитки»", title_style
        ),
        Spacer(1, 12),
        Paragraph(
            "Доставка по Москве: 300 рублей, бесплатно при заказе от "
            "4000 рублей. Срок доставки — 1-2 рабочих дня.",
            body_style,
        ),
        Spacer(1, 8),
        Paragraph(
            "Доставка по России курьерскими службами (СДЭК, Почта "
            "России): от 350 рублей в зависимости от региона. Срок "
            "доставки — от 3 до 7 рабочих дней.",
            body_style,
        ),
        Spacer(1, 8),
        Paragraph(
            "Самовывоз из шоурума в Москве — бесплатно, по "
            "предварительной договорённости с менеджером.",
            body_style,
        ),
        Spacer(1, 8),
        Paragraph(
            "Возврат товара возможен в течение 14 дней с момента "
            "получения при условии, что товар не был в носке и "
            "сохранены бирки. Если возврат оформляется по инициативе "
            "покупателя (например, не подошёл размер), доставку "
            "возврата оплачивает покупатель. Если причина возврата — "
            "производственный брак, доставку и полную стоимость "
            "товара оплачивает магазин.",
            body_style,
        ),
        Spacer(1, 8),
        Paragraph(
            "Первый обмен товара на другой размер выполняется "
            "бесплатно вне зависимости от причины обмена.",
            body_style,
        ),
    ]
    doc.build(story)


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    generate_about_txt()
    generate_faq_docx()
    generate_price_list_xlsx()
    generate_shipping_returns_pdf()
    print(f"Демо-документы созданы в {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
